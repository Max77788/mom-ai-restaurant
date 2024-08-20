import pandas as pd
from docx import Document
from fpdf import FPDF

# Text Content
menu_text = """Restaurant   MENU

1. Appetizers 
  - Garlic Bread  -    $5.99
  -Bruschetta - 7.99 $
  -    Caesar   Salad-   $8.99

2. Main     Courses
- Margherita Pizza- $12.99
     - Spaghetti Carbonara  $ 14.99
- Grilled SALMON  - $18.99

  3. DESSERTS
- Tiramisu-  $6.99
     -Chocolate   Lava  Cake- $7.99
 -Gelato- $5.99

  4. beverages
-   coffee- $2.99
-Tea -2.49 $
     -SOFT DRINK- $1.99
"""

# 1. TXT File
with open("menu_examples/menu.txt", "w") as file:
    file.write(menu_text)

# 2. XLSX File
menu_data = {
    "Category": ["Appetizers", "Appetizers", "Appetizers", "Main Courses", "Main Courses", "Main Courses", 
                 "Desserts", "Desserts", "Desserts", "Beverages", "Beverages", "Beverages"],
    "Item": ["Garlic Bread", "Bruschetta", "Caesar Salad", "Margherita Pizza", "Spaghetti Carbonara", "Grilled Salmon", 
             "Tiramisu", "Chocolate Lava Cake", "Gelato", "Coffee", "Tea", "Soft Drink"],
    "Price": ["$5.99", "$7.99", "$8.99", "$12.99", "$14.99", "$18.99", "$6.99", "$7.99", "$5.99", "$2.99", "$2.49", "$1.99"]
}

df = pd.DataFrame(menu_data)
df.to_excel("menu_examples/menu.xlsx", index=False)

# 3. DOCX File
doc = Document()
doc.add_heading('Restaurant Menu', 0)

doc.add_heading('1. Appetizers', level=1)
doc.add_paragraph('- Garlic Bread - $5.99')
doc.add_paragraph('- Bruschetta - $7.99')
doc.add_paragraph('- Caesar Salad - $8.99')

doc.add_heading('2. Main Courses', level=1)
doc.add_paragraph('- Margherita Pizza - $12.99')
doc.add_paragraph('- Spaghetti Carbonara - $14.99')
doc.add_paragraph('- Grilled Salmon - $18.99')

doc.add_heading('3. Desserts', level=1)
doc.add_paragraph('- Tiramisu - $6.99')
doc.add_paragraph('- Chocolate Lava Cake - $7.99')
doc.add_paragraph('- Gelato - $5.99')

doc.add_heading('4. Beverages', level=1)
doc.add_paragraph('- Coffee - $2.99')
doc.add_paragraph('- Tea - $2.49')
doc.add_paragraph('- Soft Drink - $1.99')

doc.save("menu_examples/menu.docx")

# 4. PDF File
pdf = FPDF()

pdf.add_page()
pdf.set_font("Arial", size=12)

# Adding text
for line in menu_text.split('\n'):
    pdf.cell(200, 10, txt=line, ln=True)

pdf.output("menu_examples/menu.pdf")

print("Files generated: menu.txt, menu.xlsx, menu.docx, menu.pdf")
